
#%%
import docx
#获取文档对象
file=docx.Document(r"C:\Users\lenovo\Desktop\fd++.docx")
print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段
 
#输出每一段的内容
for para in file.paragraphs:
    print(para.text)
    
    for i in range(len(file.paragraphs)):
        print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)


#%%



